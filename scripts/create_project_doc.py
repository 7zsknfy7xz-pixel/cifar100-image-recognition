from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "图像识别与模式分类项目说明.docx"


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    set_east_asia_font(run)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    set_east_asia_font(run)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    set_east_asia_font(run)


def build_document() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("图像识别与模式分类项目说明")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    set_east_asia_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("面向 CIFAR-10 的传统特征方法与深度学习方法对比实验")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("555555")
    set_east_asia_font(run)

    paragraph = doc.add_paragraph()
    paragraph.add_run("用途说明：").bold = True
    paragraph.add_run(
        "本文档用于整理课程大作业的项目背景、老师要求、实验方法和后续报告写作方向。"
        "它是项目说明与写作辅助材料，不直接替代最终提交的课程报告。"
    )

    doc.add_heading("1. 项目背景", level=1)
    doc.add_paragraph(
        "图像识别与模式分类是图像处理和计算机视觉中的基础任务，其目标是从图像中提取有判别力的视觉特征，"
        "并通过分类模型判断图像所属类别。传统图像处理方法通常依赖人工设计的颜色、纹理、边缘或形状特征；"
        "现代方法则更多使用卷积神经网络等深度学习模型自动学习多层视觉表示。"
    )
    doc.add_paragraph(
        "本项目选择 CIFAR-10 数据集作为实验对象。该数据集包含 10 类自然图像，图像尺寸为 32 x 32，"
        "类别包括 airplane、automobile、bird、cat、deer、dog、frog、horse、ship、truck。"
        "它规模适中、类别清晰，适合在课程大作业中比较传统机器学习方法和深度学习方法的差异。"
    )

    project_table = doc.add_table(rows=1, cols=2)
    project_table.style = "Table Grid"
    for cell, text in zip(project_table.rows[0].cells, ["项目要素", "说明"]):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, text, bold=True)
    for left, right in [
        ("研究方向", "第 9 类：其他图像与视频处理方法中的图像识别与模式分类。"),
        ("实验任务", "在 CIFAR-10 上完成多方法分类实验，并分析不同方法的性能差异。"),
        ("核心对比", "传统低维特征、人工设计特征、深度卷积特征。"),
        ("主要输出", "准确率、宏平均 Precision/Recall/F1、混淆矩阵、训练曲线、预测样例和错误样例。"),
    ]:
        row = project_table.add_row().cells
        set_cell_text(row[0], left, bold=True)
        set_cell_text(row[1], right)

    doc.add_heading("2. 老师要求整理", level=1)
    doc.add_paragraph("根据项目 PDF，本次小组作业的核心要求可以整理为以下几点：")
    for item in [
        "题目应面向图像或视频领域，可选择图像识别与模式分类方向。",
        "需要调研国内外科研团队、机构和学者的研究进展，并进行分类和分析。",
        "需要介绍图像处理方法的原理和过程，建议给出框图和公式。",
        "至少完成三个方法的实验，说明实验设置，展示不同图像或参数配置下的结果，并分析性能。",
        "需要给出图表形式的实验分析。",
        "需要思考和评论该技术未来发展趋势。",
        "参考文献不少于 20 篇，近 10 年文献占 50% 以上，并合理引用。",
        "报告篇幅为 8-10 页，包含参考文献，A4 格式。",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph(
        "评分标准中，撰写格式约占 20%，内容约占 50%，实验约占 30%。"
        "因此，本项目代码重点服务于实验部分，同时为报告中的方法解释和图表分析提供支撑。"
    )

    doc.add_heading("3. 方法路线", level=1)
    methods_table = doc.add_table(rows=1, cols=4)
    methods_table.style = "Table Grid"
    for cell, text in zip(methods_table.rows[0].cells, ["方法", "特征来源", "分类器/模型", "对比意义"]):
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True)
    for row_data in [
        ("PCA + SVM", "灰度像素展平后 PCA 降维", "LinearSVM", "作为低维统计特征和传统分类器基线。"),
        ("HOG/颜色特征 + SVM", "梯度方向直方图、RGB/HSV 颜色直方图", "LinearSVM", "体现人工设计图像特征在分类任务中的作用。"),
        ("ResNet18", "卷积网络自动学习多层视觉特征", "端到端神经网络", "代表深度学习方法，通常具备更强表达能力。"),
    ]:
        row = methods_table.add_row().cells
        for idx, text in enumerate(row_data):
            set_cell_text(row[idx], text, bold=(idx == 0))

    doc.add_heading("3.1 PCA + SVM", level=2)
    doc.add_paragraph(
        "该方法先将彩色图像转换为灰度图并展平为向量，再用 PCA 将高维像素空间投影到低维主成分空间。"
        "PCA 保留数据中方差较大的方向，减少冗余和噪声，最后使用线性 SVM 进行分类。"
    )

    doc.add_heading("3.2 HOG/颜色特征 + SVM", level=2)
    doc.add_paragraph(
        "HOG 描述局部梯度方向分布，适合表达边缘、轮廓和局部形状；颜色直方图描述 RGB/HSV 空间中的颜色分布。"
        "二者拼接后输入线性 SVM，可以比较人工特征相对于原始像素特征的提升。"
    )

    doc.add_heading("3.3 ResNet18", level=2)
    doc.add_paragraph(
        "ResNet18 使用残差连接缓解深层网络训练中的退化问题。项目代码将第一层卷积和池化结构改为更适合 "
        "32 x 32 CIFAR-10 图像的形式，并使用交叉熵损失、AdamW 优化器和余弦学习率调度进行训练。"
    )

    doc.add_heading("4. 实验设计", level=1)
    for item in [
        "数据集：CIFAR-10，训练集 50000 张，测试集 10000 张，10 个类别。",
        "评价指标：Accuracy、Macro Precision、Macro Recall、Macro F1、每类准确率。",
        "可视化结果：混淆矩阵、训练曲线、预测样例、错误分类样例。",
        "调试方式：先使用小样本和 1 个 epoch 验证流程，再运行完整实验。",
        "硬件环境：WSL Ubuntu 20.04，优先使用 RTX 4060 GPU 训练 ResNet18。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. 后续报告写作建议", level=1)
    for item in [
        "背景部分可说明图像识别从人工特征到深度学习的发展脉络。",
        "方法部分建议为三种方法分别给出流程图、核心公式和适用特点。",
        "实验部分可先给出统一设置，再展示指标表和混淆矩阵。",
        "分析部分重点讨论传统方法和深度模型在类别区分能力、错误样例和计算成本上的差异。",
        "趋势部分可讨论自监督学习、视觉 Transformer、多模态模型、轻量化部署和可信 AI。",
        "最终报告需要小组成员结合实验结果自行核对、改写和补充引用，避免直接提交工具生成文本。",
    ]:
        add_number(doc, item)

    doc.add_heading("6. 建议运行步骤", level=1)
    for item in [
        "在 VS Code 中使用 Reopen Folder in WSL 打开项目目录。",
        "创建并激活 Python 虚拟环境。",
        "执行 pip install -r requirements.txt 安装依赖。",
        "运行 GPU 检查命令确认 torch.cuda.is_available() 为 True。",
        "先运行小样本测试命令，确认 outputs 目录下有指标表和图片。",
        "最后运行正式实验命令，并将 outputs 中的图表用于报告分析。",
    ]:
        add_number(doc, item)

    doc.add_paragraph()
    end = doc.add_paragraph()
    run = end.add_run("备注：本说明文档根据课程 PDF 要求和当前实验项目设计整理，供课程大作业实施和报告写作参考。")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")
    set_east_asia_font(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("图像处理课程大作业项目说明")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("777777")
    set_east_asia_font(run)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
